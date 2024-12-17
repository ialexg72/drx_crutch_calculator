from lxml import etree
from typing import List
import os
import uuid
import shutil
import subprocess
import docx
from docx.table import Table, _Row, _Cell
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import logging
import logging.config
from src import settings, utility
from docx.shared import Inches
# Импортируем настройки логирования
logging.config.dictConfig(settings.LOGGING_CONFIG)
logger = logging.getLogger(__name__)

def load_drawio_file_lxml(file_path: str) -> etree._ElementTree:
        try:
            parser = etree.XMLParser(remove_comments=False)
            tree = etree.parse(file_path, parser)
            return tree
        except etree.XMLSyntaxError as e:
            raise ValueError(f"Ошибка парсинга XML файла: {e}")
        except FileNotFoundError:
            raise ValueError("Файл не найден. Проверьте путь к файлу.")

def find_layers(tree: etree._ElementTree, layer_names: List[str]) -> List[etree._Element]:
    """
    Возвращает список слоев по их названиям.
    """
    layers = []
    # XPath для поиска <mxCell> с parent="0" и value равным одному из имён слоёв
    xpath_query = ".//mxCell[@parent='0' and ("
    xpath_conditions = []
    for name in layer_names:
        xpath_conditions.append(f"@value='{name}'")
    xpath_query += " or ".join(xpath_conditions) + ")]"
    
    layers = tree.xpath(xpath_query)
    logger.debug(f"Найдено слоев: {len(layers)}, искомые слои: {layer_names}")
    return layers

def toggle_layer_visibility(tree: etree._ElementTree, layers: List[etree._Element], visibility: bool) -> None:
    """
    Устанавливает видимость для указанных слоёв.
    """
    for layer in layers:
        # Установка атрибута 'visible'
        layer.set("visible", "1" if visibility else "0")
        layer_name = layer.get('value')
        logger.info(f"Слой '{layer_name}' установлен {('видимым' if visibility else 'невидимым')}")

def save_drawio_as_png(tree: etree._ElementTree, scheme_template_path: str, organization: str, save_dir: str = "tmp") -> str:
    temp_drawio_path, _ = utility.generate_filename(organization, "drawio")
    png_output_path, _ = utility.generate_filename(organization, "png")
    
    # Запись временного файла
    try:
        tree.write(temp_drawio_path, pretty_print=True, xml_declaration=True, encoding='UTF-8')
        logger.debug(f"Временный .drawio файл создан по пути: {temp_drawio_path}")
    except Exception as e:
        logger.error(f"Не удалось записать временный файл: {e}")
        raise
    
    # Указание пути к исполняемому файлу drawio-exporter
    #drawio_exporter_executable = r"C:\Program Files\draw.io\draw.io.exe"
    drawio_exporter_executable = r"drawio"
    
    # Проверка наличия исполняемого файла в PATH или по указанному пути
    if not shutil.which(drawio_exporter_executable):
        raise FileNotFoundError(f"Исполняемый файл drawio-exporter не найден. Убедитесь, что он установлен и доступен в PATH.")
    
    command = [
        drawio_exporter_executable,
        '-x', temp_drawio_path,
        '-o', png_output_path,
        '-f', 'png',
        '-b', '5',
        '--no-sandbox'
    ]
    
    logger.debug(f"Выполнение команды: {' '.join(command)}")
    
    try:
        result = subprocess.run(command, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, encoding='utf-8')
        logger.debug(f"drawio-exporter успешно конвертировал файл. Вывод: {result.stdout}")
        if result.stderr:
            logger.warning(f"Предупреждение от drawio-exporter: {result.stderr}")
    except subprocess.CalledProcessError as e:
        logger.error(f"Ошибка при конвертации в PNG: {e.stderr}")
        raise RuntimeError(f"Ошибка при конвертации в PNG: {e.stderr.strip()}") from e
    
    # Проверка существования PNG-файла
    if not os.path.isfile(png_output_path):
        logger.error(f"PNG-файл не был создан по пути: {png_output_path}")
        raise FileNotFoundError(f"PNG-файл не был создан по пути: {png_output_path}")
    
    # Удаление временного файла
    try:
        os.remove(temp_drawio_path)
        logger.debug(f"Временный файл {temp_drawio_path} удален.")
    except OSError as e:
        logger.error(f"Не удалось удалить временный файл {temp_drawio_path}: {e}")
    
    return png_output_path

def replace_placeholder_with_image(doc, placeholder, image_path, width_inches=None):
    """
    Заменяет указанный текст-заполнитель на изображение в документе Word.
    """       
    if not os.path.exists(image_path):
        logger.error(f"Изображение не найдено: {image_path}")
        raise ValueError(f"Изображение не найдено: {image_path}")

    replaced = False

    def replace_in_paragraphs(paragraphs):
        nonlocal replaced
        for paragraph in paragraphs:
            if placeholder in paragraph.text:
                # Объединение всех runs в одном тексте
                inline = paragraph.runs
                for i in range(len(inline)):
                    if placeholder in inline[i].text:
                        text = inline[i].text.replace(placeholder, "")
                        inline[i].text = text
                        # Добавляем изображение только один раз после замены placeholder
                        run = paragraph.add_run()
                        try:
                            if width_inches:
                                run.add_picture(image_path, width=Inches(width_inches))
                                logger.info(f"Вставлено изображение '{image_path}' с шириной {width_inches} дюймов.")
                            else:
                                run.add_picture(image_path)
                                logger.info(f"Вставлено изображение '{image_path}' без указания ширины.")
                            replaced = True
                        except Exception as e:
                            raise ValueError(f"Ошибка при вставке изображения: {e}")

    # Обработка основных параграфов
    replace_in_paragraphs(doc.paragraphs)

    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)

    # Обработка заголовков и нижних колонтитулов
    for section in doc.sections:
        replace_in_paragraphs(section.header.paragraphs)
        replace_in_paragraphs(section.footer.paragraphs)

    if not replaced:
        logger.warning(f"Заполнитель '{placeholder}' не найден в документе.")
        raise ValueError(f"Заполнитель '{placeholder}' не найден в документе.")

    return replaced

def drawing_scheme(redundancy, layers_to_toggle, template_path, scheme_template, organization):  
    visibility = False
    logger.info(f"Начало обработки схемы для организации: {organization}")
    logger.info(f"Слои для переключения: {layers_to_toggle}")

    # Шаг 1: Загрузка и парсинг файла
    try:
        tree = load_drawio_file_lxml(scheme_template)
        logger.info(f"Файл схемы успешно загружен: {scheme_template}")
    except ValueError as e:
        logger.error(f"Ошибка при загрузке файла схемы: {e}")
        raise

    if layers_to_toggle:
        try:
            # Шаг 2: Поиск слоёв
            layers = find_layers(tree, layers_to_toggle)
            if not layers:
                logger.warning("Не найдено ни одного слоя из списка для переключения")
            
            # Шаг 3: Изменение видимости слоёв
            toggle_layer_visibility(tree, layers, visibility)
            logger.info(f"Видимость слоев успешно изменена. Всего обработано слоев: {len(layers)}")
        except Exception as e:
            logger.error(f"Ошибка при обработке слоев: {e}")
            raise
    else:
        logger.info("Список слоев для переключения пуст, пропускаем обработку слоев")

    # Шаг 4: Сохранение файла
    try:
        saved_file = save_drawio_as_png(tree, scheme_template, organization)
        logger.info(f"Схема успешно сохранена в файл: {saved_file}")
        return saved_file
    except Exception as e:
        logger.error(f"Ошибка при сохранении схемы: {e}")
        raise