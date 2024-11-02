# app.py
import os
import uuid
import math
import docx
import xml.etree.ElementTree as ET
from flask import Flask, request, render_template, send_from_directory, redirect, url_for
from docx.table import _Row
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime

app = Flask(__name__)

# Папки для загрузок и отчетов
UPLOAD_FOLDER = 'uploads'
REPORT_FOLDER = 'ready_reports'
TEMPLATE_FOLDER = 'word_templates'

# Убедимся, что папки существуют
for folder in [UPLOAD_FOLDER, REPORT_FOLDER, TEMPLATE_FOLDER]:
    if not os.path.exists(folder):
        os.makedirs(folder)

# Настройки Flask
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['REPORT_FOLDER'] = REPORT_FOLDER
app.config['TEMPLATE_FOLDER'] = TEMPLATE_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # Максимальный размер файла: 16MB

@app.route('/', methods=['GET'])
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_xml():
    if 'xml_file' not in request.files:
        return "Файл не загружен", 400
    file = request.files['xml_file']
    if file.filename == '':
        return "Имя файла пустое", 400
    if not file.filename.lower().endswith('.xml'):
        return "Неподдерживаемый формат файла. Пожалуйста, загрузите XML файл.", 400

    # Сохраняем загруженный XML файл
    filename = f"{uuid.uuid4()}.xml"
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    # Парсим XML файл
    try:
        tree = ET.parse(filepath)
        root = tree.getroot()
    except ET.ParseError:
        return "Ошибка при разборе XML файла.", 400

    # Извлекаем необходимые данные из XML
    data = {}
    for child in root:
        data[child.tag] = child.text

    #Функция получения даты
    def get_current_date_formatted():
        today = datetime.today()
        return today.strftime("%d.%m.%Y")
    current_date = get_current_date_formatted()

    # Закидываем данные из XML в переменные
        #Общая информация
    organization = data.get('organization', '')
    operationsystem = data.get('ostype', '')
    version = data.get('version', '')
    kubernetes = data.get('kubernetes', '')
    redundancy = data.get('redundancy', '')
    monitoring = data.get('monitoring', '')
    database = data.get('database', '')
        #Активность пользователей
    registeredUsers = int(data.get('registeredUsers', ''))
    peakLoad = int(data.get('peakLoad', ''))
    peakPeriod = data.get('peakPeriod', '')
    concurrent_users = int(data.get('concurrentUsers', ''))
    mobileusers = int(data.get('mobileappusers', ''))
    lk_users = int(data.get('lkusers', ''))
        #Прирост данных
    importhistorydata = int(data.get('importhistorydata', ''))
    annualdatagrowth = int(data.get('annualdatagrowth', ''))
    midsizedoc = int(data.get('midsizedoc', ''))
        #Импорт данных в систему
    dcs = data.get('dcs', '')
    dcsdochours = int(data.get('dcsdochours', ''))
        #Интеграция
    onlineeditor = data.get('onlineeditor', '')
    integrationsystems = data.get('integrationsystems', '')
        #Поиск и обработка данных
    elasticsearch = data.get('elasticsearch', '')
    ario = data.get('ario', '')
    ariodocin = int(data.get('ariodocin', ''))


    # Загружаем шаблон Word
    if operationsystem.lower() == "linux":
        template_path = os.path.join(app.config['TEMPLATE_FOLDER'], 'RecomendBaseTpl4.10_linux.docx')
    else:
        template_path = os.path.join(app.config['TEMPLATE_FOLDER'], 'RecomendBaseTpl4.10_winux.docx')
    if not os.path.exists(template_path):
        print("Шаблон Word не найден.", 500)
    doc = docx.Document(template_path)

    #Функция для удаления не нужных блоков в таблицах
    def delete_row_from_table(table, row):
           tbl = table._tbl
           tr = row._tr
           tbl.remove(tr)

    def remove_specific_rows(doc_path, target_text, num_rows_to_delete=5):
       # Открываем документ 
       # Проходим по всем таблицам в документе
       for table in doc.tables:
           # Индекс текущей строки
           i = 0
           while i < len(table.rows):
               row = table.rows[i]
               # Проверяем все ячейки в строке на наличие целевого текста
               if any(target_text in cell.text for cell in row.cells):
                   # Удаляем найденную строку и следующие num_rows_to_delete строк
                   for _ in range(num_rows_to_delete + 1):  # +1 для самой найденной строки
                       if i < len(table.rows):
                           delete_row_from_table(table, table.rows[i])
                       else:
                           break
                   # После удаления сдвигаем индекс обратно, чтобы продолжить проверку
                   i -= 1
               i += 1

    #Функция поиска и удаления текста
    def delete_paragraphs_by_text(doc, text_to_delete):
        paragraphs = doc.paragraphs
        for paragraph in paragraphs:
            if text_to_delete in paragraph.text:
                p = paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None

    # Расчеты Kubernetes Control-plane
    if kubernetes.lower() == "true":
        if redundancy.lower() == "true":
            k8s_count = 3
        else:
            k8s_count = 1
        k8s_cpu = 4
        k8s_ram = 4
        k8s_hdd = 50
    else:
        k8s_count = 0
        k8s_cpu = 0
        k8s_ram = 0
        k8s_hdd = 0
        if len(doc.tables) > 4:
            table = doc.tables[4]
            start_index = 1
            end_index = 7
            remove_specific_rows(doc, "Узел администрирования Kubernetes", 6)
            delete_paragraphs_by_text(doc, "Kubernetes")
            delete_paragraphs_by_text(doc, "На узле генерируется конфигурационный файл config.yml")
        else:
            print("Error")

    #Расчет узлов веб-серверов
    def calculate_webserver_count(concurrent_users, redundancy):
        if concurrent_users > 0:
            count = math.ceil(concurrent_users / 2500)
            return count + 1 if redundancy.lower() == "true" else count
        return 0
    def round_up_to_even(value):
        return value if value % 2 == 0 else value + 1
    def compute_srvunit(concurrent_users, webserver_count, redundancy):
        if concurrent_users == 0 or webserver_count == 0:
            srv_cpu = 0
        elif concurrent_users < 501:
            srv_cpu = 6
        else:
            divider = webserver_count - 1 if redundancy.lower() == "true" else webserver_count
            if redundancy.lower() == "true" and webserver_count <= 1:
                raise ValueError("При redundancy='true' значение webserver_count должно быть больше 1.")
            temp = math.ceil(concurrent_users / divider / 500)
            srv_cpu = temp * 2 + 2
        return round_up_to_even(srv_cpu)
    def calculate_srvunitram_ws(concurrent_users, webserver_count, redundancy):
        if concurrent_users == 0:
            temp = 0
        elif concurrent_users < 501:
            temp = 14
        elif concurrent_users < 2501:
            temp = 12
        else:
            divider = webserver_count - 1 if redundancy.lower() == "true" else webserver_count
            if redundancy.lower() == "true" and webserver_count <= 1:
                raise ValueError("webserver_count должно быть больше 1, если redundancy равно 'Да'.")
            ceil_value = math.ceil(concurrent_users / divider / 500)
            temp = ceil_value * 2 + 2
        return temp if temp % 2 == 0 else temp + 1
    webserver_count = calculate_webserver_count(concurrent_users, redundancy)
    webserver_cpu = compute_srvunit(concurrent_users, webserver_count, redundancy)
    webserver_ram = calculate_srvunitram_ws(concurrent_users, webserver_count, redundancy)
    webserver_hdd = 100 if webserver_cpu != 0 else 0

    # Расчет узлов микросервисов
    if concurrent_users > 500:
        def calculate_ms_count(concurrent_users, redundancy):
            if concurrent_users > 500:
                ms_count = math.ceil(concurrent_users / 2500)
                return ms_count + 1 if redundancy.lower() == "true" else ms_count
            return 0
        def round_up_to_even(value):
            return value if value % 2 == 0 else value + 1
        def calculate_ms_cpu(concurrent_users, ms_count, redundancy):
            if concurrent_users < 1001:
                result = 6
            elif redundancy.lower() == "true":
                result = math.ceil(concurrent_users / (ms_count - 1) / 500) * 2
            else:
                result = math.ceil(concurrent_users / ms_count / 500) * 2
            return result
        def calculate_ms_ram(concurrent_users, ms_count, redundancy):
            if concurrent_users == 0 or ms_count == 0:
                return 0
            if concurrent_users < 1501:
                return 12
            divider = ms_count - 1 if redundancy.lower() == "true" else ms_count
            if redundancy.lower() == "true" and ms_count <= 1:
                raise ValueError("webserver_count должно быть больше 1, если redundancy равно 'Да'.")
            temp = math.ceil(concurrent_users / divider / 1000)
            value = temp * 6
            return value if value % 2 == 0 else value + 1
        def calculate_ms_hdd(concurrent_users):
            return 100 if concurrent_users > 500 else 0
        ms_count = calculate_ms_count(concurrent_users, redundancy)
        ms_cpu = calculate_ms_cpu(concurrent_users, ms_count, redundancy)
        ms_ram = calculate_ms_ram(concurrent_users, ms_count, redundancy)
        ms_hdd = calculate_ms_hdd(concurrent_users)
    else:
        ms_count = 0
        ms_cpu = 0
        ms_ram = 0
        ms_hdd = 0
        remove_specific_rows(doc, "Узлы микросервисов Directum RX", 6)
        delete_paragraphs_by_text(doc, "Узлы микросервисов Directum RX")
        delete_paragraphs_by_text(doc, "На узле генерируется конфигурационный файл config.yml")

    #Расчеты для сервиса Nomad
    if mobileusers != 0:
        #Расчет ЦПУ
        def calculate_nomad_count(mobileusers, redundancy):
            if redundancy.lower() == "true":
                result = math.ceil(mobileusers / 1000) + 1
            else:
                result = math.ceil(mobileusers / 1000)
            return result
        nomad_count = calculate_nomad_count(mobileusers, redundancy)
        def calculate_nomad_cpu(mobileusers, redundancy, nomad_count):
            if redundancy.lower() == "true":
                temp_result = math.ceil(mobileusers / (nomad_count - 1) / 150) * 2 + 2
            else:
                temp_result = math.ceil(mobileusers / nomad_count / 150) * 2 + 2  
            result = round_up_to_even(temp_result)
            return result
        nomad_cpu = calculate_nomad_cpu(mobileusers, redundancy, nomad_count)
        def calculate_nomad_ram(mobileusers, redundancy, nomad_count):
            if redundancy.lower() == "true":
                result = math.ceil(mobileusers / (nomad_count - 1) / 50 * 1.5 + 2)
            else:
                result = math.ceil(mobileusers / nomad_count / 50 * 1.5 + 2)
            if round_up_to_even(result):
                return result
            else:
                return result + 1
        nomad_ram = calculate_nomad_ram(mobileusers, redundancy, nomad_count)
        nomad_hdd = 100
    else:
        nomad_count = 0
        nomad_cpu = 0
        nomad_ram = 0
        nomad_hdd = 0
        remove_specific_rows(doc, "Узлы сервиса NOMAD", 6)
        delete_paragraphs_by_text(doc, "Сервис NOMAD (NomadService)")

    #4 Считаем reverse proxy узлы
    if concurrent_users > 500 or redundancy.lower() == "true":
        def calculate_reverseproxy_count(concurrent_users, redundancy):
            if concurrent_users != 0:
                if redundancy.lower() == "true":
                    result = 2
                else:
                    if concurrent_users < 500:
                        result = 0
                    else:
                        result = 1
            else:
                result = 0
            return result
        reverseproxy_count = calculate_reverseproxy_count(concurrent_users, redundancy)
        def calculate_reverseproxy_cpu(reverseproxy_count, concurrent_users):
            if reverseproxy_count != 0:
                divided = concurrent_users / 5000
                ceiled = math.ceil(divided)
                multiplied = ceiled * 2
            else:
                multiplied = 0
            if multiplied % 2 != 0:
                multiplied += 1 
            return multiplied
        reverseproxy_cpu = calculate_reverseproxy_cpu(reverseproxy_count, concurrent_users)
        def calculate_reverseproxy_ram(reverseproxy_count, concurrent_users):
            if reverseproxy_count != 0:
                divided = concurrent_users / 5000
                ceiled = math.ceil(divided)
                multiplied = ceiled * 2
            else:
                multiplied = 0
            if multiplied % 2 != 0:
                multiplied += 1 
            return multiplied
        reverseproxy_ram = calculate_reverseproxy_ram(reverseproxy_count, concurrent_users)
        def calculate_reverseproxy_hdd(reverseproxy_count):
            if reverseproxy_count != 0:
                rp_hdd = 50
            else:
                rp_hdd = 0
            return rp_hdd
        reverseproxy_hdd = calculate_reverseproxy_hdd(reverseproxy_count)
    else:
        reverseproxy_count = 0
        reverseproxy_cpu = 0
        reverseproxy_ram = 0 
        reverseproxy_hdd = 0
        remove_specific_rows(doc, "Узлы reverse proxy", 6)
        delete_paragraphs_by_text(doc, "reverse-proxy")
    #5 Расчет СУБД
    def calculate_sql_count(redundancy):
        if redundancy.lower() == "true":
            count = 2
        else:
            count = 1
        return count
    sql_count = calculate_sql_count(redundancy)

    def calculate_sql_cpu(concurrent_users, redundancy, lk_users):
        if concurrent_users < 501:
            cpu = 6 + (lk_users/10000)*2
        elif concurrent_users < 1500:
            cpu = 8 + (lk_users/10000)*2
        else:
            cpu = math.ceil(concurrent_users/400)*2+(lk_users/10000)*2
        if cpu >= 0:
            ceil_num = math.ceil(cpu)
            if ceil_num % 2 == 0:
                return ceil_num
            else:
                return ceil_num + 1
        else:
            floor_num = math.floor(number)
            if floor_num % 2 == 0:
                return floor_num
            else:
                return floor_num - 1
    sql_cpu = calculate_sql_cpu(concurrent_users, redundancy, lk_users)

    def calculate_sql_ram(concurrent_users, redundancy, lk_users):
        if concurrent_users > 50 or redundancy == "Да":
            if concurrent_users < 500:
                value = math.ceil(concurrent_users / 125) + 6 + (lk_users / 10000) * 4
            elif concurrent_users < 2000:
                value = 16 + (lk_users / 10000) * 4
            else:
                value = math.ceil(concurrent_users / 400) * 4 + (lk_users / 10000) * 4
        else:
            value = 0
        rounded_value = math.ceil(value)
        if rounded_value % 2 != 0:
            even_value = rounded_value + 1
        else:
            even_value = rounded_value
        return even_value
    sql_ram = calculate_sql_ram(concurrent_users, redundancy, lk_users)

    #Служба ввода документов DCS
    if dcs.lower() == "true":
        dcs_count = 1
        def calculate_dcs_cpu(dcs, dcsdochours):
            rounded_up = math.ceil(dcsdochours / 150)
            intermediate_result = rounded_up + 2
            if intermediate_result % 2 != 0:
                final_result = intermediate_result + 1
            else:
                final_result = intermediate_result
            return final_result
        dcs_cpu = calculate_dcs_cpu(dcs, dcsdochours)
        def calculate_dcs_ram(dcs, dcsdochours):
            rounded_up = math.ceil(dcsdochours / 150)*2
            intermediate_result = rounded_up + 2
            if intermediate_result % 2 != 0:
                final_result = intermediate_result + 1
            else:
                final_result = intermediate_result
            return final_result
        dcs_ram = calculate_dcs_ram(dcs, dcsdochours)
        dcs_hdd = 50
    else:
        dcs_count = 0
        dcs_cpu = 0
        dcs_ram = 0
        dcs_hdd = 0
        remove_specific_rows(doc, "Узел службы ввода документов", 6)
        delete_paragraphs_by_text(doc, "Узлы DCS")

    # Полнотекстовый поиск
    if elasticsearch.lower() == "true":
        elasticsearch_count = 1
        elasticsearch_cpu = 8
        def calculate_search_ram(elasticsearch, midsizedoc, annualdatagrowth):
            ram = annualdatagrowth * midsizedoc
            gigabytes = ram / (1024 ** 3)
            if gigabytes > 6:
                intermediate_result = 32
            else:
                intermediate_result = 16
            if intermediate_result % 2 == 0:
               return intermediate_result
            elif intermediate_result > 0:
               return intermediate_result + 1
            else:
               return intermediate_result - 1
        elasticsearch_ram = calculate_search_ram(elasticsearch, midsizedoc, annualdatagrowth)
        elasticsearch_hdd = 50
    else:
        elasticsearch_count = 0
        elasticsearch_cpu = 0
        elasticsearch_ram = 0
        elasticsearch_hdd = 0
        remove_specific_rows(doc, "Узел полнотекстового поиска", 6)
        remove_specific_rows(doc, "Разделы для индексов полнотекстового поиска", 1)
        delete_paragraphs_by_text(doc, "Узел полнотекстового поиска – виртуальная машина")
        delete_paragraphs_by_text(doc, "Хранилище для индексов полнотекстового поиска")

    #Мониторинг
    if monitoring.lower() == "true":
        monitoring_count = 1
        monitoring_hdd = 50
        monitoring_cpu = 16 if concurrent_users > 3000 else 8
        monitoring_ram = 32 if concurrent_users > 3000 else 16 
        if concurrent_users > 2000:
            logstash_count = 1
            logstash_cpu = 4
            logstash_ram = 6
            logstash_hdd = 50
        else: 
            logstash_count = 0
            logstash_cpu = 0
            logstash_ram = 0
            logstash_hdd = 0
            remove_specific_rows(doc, "Узел Logstash", 6)
    else:
        monitoring_count = 0
        monitoring_hdd = 0
        monitoring_cpu = 0
        monitoring_ram = 0
        logstash_count = 0
        logstash_cpu = 0
        logstash_ram = 0
        logstash_hdd = 0
        remove_specific_rows(doc, "Узел решения «Мониторинг системы Directum RX»", 6)
        delete_paragraphs_by_text(doc, "Узел решения «Мониторинг системы Directum RX»")
        remove_specific_rows(doc, "Узел Logstash", 6)
        remove_specific_rows(doc, "Разделы для индексов системы мониторинга", 1)

    #Узлы АРИО
    if ario.lower() == "true":
        if operationsystem.lower() == "linux":    
            ario_count = 1
            ario_hdd = 100
            def calculate_ario_cpu(ariodocin):
                if ariodocin <= 25000:
                    return 4
                elif ariodocin <= 55000:
                    return 8
                elif ariodocin <= 90000:
                    return 12
                elif ariodocin <= 150000:
                    return 10
                elif ariodocin <= 250000:
                    return 16
                else:
                    return "Error"           
            ario_cpu = calculate_ario_cpu(ariodocin)

            def calculate_ario_ram(ariodocin):
                if ariodocin <= 25000:
                    return 20
                elif ariodocin <= 55000:
                    return 24
                elif ariodocin <= 90000:
                    return 40
                elif ariodocin <= 150000:
                    return 14
                elif ariodocin <= 250000:
                    return 24
                else:
                    return "Error"
            ario_ram = calculate_ario_ram(ariodocin)

            if ariodocin > 90000:
                dtes_count = 1
                dtes_hdd = 100
                def calculate_dtes_cpu(ariodocin):
                    if ariodocin <= 150000:
                        return 10
                    elif ariodocin <= 250000:
                        return 16
                    else:
                        return "Error"
                dtes_cpu = calculate_dtes_cpu(ariodocin)
                def calculate_dtes_ram(ariodocin):
                    if ariodocin <= 150000:
                        return 28
                    elif ariodocin <= 250000:
                        return 48
                    else:
                        return "Error"
                dtes_ram = calculate_dtes_ram(ariodocin)
            else:
                dtes_count = 0
                dtes_cpu = 0
                dtes_ram = 0
                dtes_hdd = 0
                remove_specific_rows(doc, "Узел сервисов Directum Text Extractor Service", 6)
    else:
        dtes_count = 0
        dtes_cpu = 0
        dtes_ram = 0
        dtes_hdd = 0
        ario_count = 0
        ario_cpu = 0
        ario_ram = 0
        ario_hdd = 0
        remove_specific_rows(doc, "Узел сервисов Directum Ario", 6)
        remove_specific_rows(doc, "Узел сервисов Directum Text Extractor Service", 6)
        remove_specific_rows(doc, "Сервисы Ario", 1)
        delete_paragraphs_by_text(doc, "Сервисы Ario")
        delete_paragraphs_by_text(doc, "** - для сервисов Ario рекомендуется использовать процессоры")

    #Узлы RRM
    if redundancy.lower() == "true":
        rrm_count = 3
        rrm_hdd = 50  
    else:
        rrm_count = 1
        rrm_hdd = 50
    if concurrent_users < 5000:
        rrm_cpu = rrm_ram = 2 
        rrm_hdd = 50
    elif concurrent_users > 10000:
        rrm_cpu = rrm_ram = 6 
        rrm_hdd = 50
    else: 
        rrm_cpu = rrm_ram = 4
        rrm_hdd = 50
    if concurrent_users > 500:
        if concurrent_users < 5000:
            rrm_cpu = rrm_ram = 2 
            rrm_hdd = 50
        elif concurrent_users > 10000:
            rrm_cpu = rrm_ram = 6 
            rrm_hdd = 50
        else: 
            rrm_cpu = rrm_ram = 4
            rrm_hdd = 50
    else:
        rrm_count = 0
        rrm_cpu = 0
        rrm_ram = 0
        rrm_hdd = 0
        delete_paragraphs_by_text(doc, "Узлы RabbitMQ, etcd+haproxy+keepalived (RMQ + EHK)")
        remove_specific_rows(doc, "Узлы RabbitMQ, etcd + keepalived + haproxy (для кластера PG)", 6)



    #Узлы интеграции с онлайн редакторами
    if onlineeditor != "none":
        onlineeditor_count = 1
        onlineeditor_hdd = 50
        def calculate_onlineeditor_cpu(concurrent_users):
            value = 2 if math.ceil(concurrent_users * 0.2) < 200 else math.floor((concurrent_users * 0.2) / 200) * 2
            return value if value % 2 == 0 else value + 1
        onlineeditor_cpu = calculate_onlineeditor_cpu(concurrent_users)
        def calculate_onlineeditor_ram(concurrent_users):
            value = 4 if math.ceil(concurrent_users * 0.2) < 200 else math.floor((concurrent_users * 0.2) / 200) * 2 + 2
            return value if value % 2 == 0 else value + 1
        onlineeditor_ram = calculate_onlineeditor_ram(concurrent_users)
    else:
        onlineeditor_count = 0
        onlineeditor_cpu = 0
        onlineeditor_ram = 0
        onlineeditor_hdd = 0
        remove_specific_rows(doc, "Узел решения «Интеграция с онлайн-редакторами OnlyOffice и Р7-Офис»", 6)
        delete_paragraphs_by_text(doc, "Узел решения «Интеграция с онлайн-редакторами»")

    #Личны кабинет
    if lk_users != 0:
        #кол-во узлов
        lk_hdd = 50
        if redundancy.lower() == "true" or concurrent_users > 5000:
            lk_count = 3
        elif concurrent_users > 75000:
            lk_count = 5
        else:
            lk_count = 1
        #ЦПУ
        if lk_count == 1:
            lk_cpu = 6
        else:
            if lk_users < 50000:
                lk_cpu = 4
            else:
                lk_cpu = 6
        lk_cpu = (lk_cpu % 2 == 0)
        #RAM
        if lk_count == 1:
            if lk_users < 1000:
                lk_ram = 12
            else:
                lk_ram = 18
        else:
            if lk_users < 50000:
                lk_ram = 8
            else:
                lk_ram = 12
        lk_ram = (lk_ram % 2 == 0)
    else:
        lk_count = 0
        lk_cpu = 0
        lk_ram = 0
        lk_hdd = 0
        delete_paragraphs_by_text(doc, "«Личный кабинет» - решение позволяет")
        delete_paragraphs_by_text(doc, "Архитектура платформы личного кабинета")
        delete_paragraphs_by_text(doc, "Сервер приложения личного кабинета")
        delete_paragraphs_by_text(doc, "Сайт личного кабинета (EssSite)")
        delete_paragraphs_by_text(doc, "Сервис идентификации (IdentityService)")
        delete_paragraphs_by_text(doc, "Cервис подписания (SignService)")
        delete_paragraphs_by_text(doc, "Сервис документов (DocumentService)")
        delete_paragraphs_by_text(doc, "Сервис сообщений (MessageBroker)")
        delete_paragraphs_by_text(doc, "Cервис предпросмотра (PreviewService)")
        delete_paragraphs_by_text(doc, "Сервис хранения файлов предпросмотра (PreviewStorage)")
        delete_paragraphs_by_text(doc, "Сервис хранения BLOB-объектов (BlobStorageService)")
        delete_paragraphs_by_text(doc, "Сервер размещения контента (ContentServer)")
        delete_paragraphs_by_text(doc, "Сервер сеансов (SessionServer)")
        remove_specific_rows(doc, "Узлы решения «Личный кабинет»", 6)
        remove_specific_rows(doc, "Узел сервисов решения «Личный кабинет»", 6)
        remove_specific_rows(doc, "HR Pro (личный кабинет)", 1)

    #Расчет хранилищ
    #Исторические данные
    importhistorydata_size = importhistorydata * midsizedoc /1024 / 1024
    importhistorydata_size = math.ceil(importhistorydata_size)
    #Годовой прирост документов
    annualdatagrowth_size = annualdatagrowth * midsizedoc / 1024 / 1024
    annualdatagrowth_size = math.ceil(annualdatagrowth_size) 
    #Объем основого хранилища тел документов
    main_storage_doc = math.ceil((annualdatagrowth_size * 6) + importhistorydata_size)
    #Объем резервного хранилища
    reserve_storage_doc = math.ceil(main_storage_doc*2)
    #Объем основного хранилища БД
    if concurrent_users != 0 or sql_count != 0:
        main_storage_db = main_storage_doc * 0.025 + (concurrent_users / 100 * 5)
        if redundancy.lower() == "true" and database.lower() == "postgres":
            if main_storage_db < 100:
                main_storage_db = 200
            else:
                main_storage_db = main_storage_db * 2
        else:
            if main_storage_db < 100:
                main_storage_db = 100
            else:
                main_storage_db  
        main_storage_db = math.ceil(main_storage_db)
    else:
        main_storage_db = 0
    #Объем резервного хранилища БД
    if database.lower() == "postgres" and redundancy.lower() == "true":
        reserve_storage_db = main_storage_db*8/2
    else:
        reserve_storage_db = main_storage_db*8
    #Разделы высокоскоростных данных (разделы ВМ под ОС, разделы БД)
    highspeed_storage = (
        main_storage_db + webserver_hdd + ms_hdd + reverseproxy_hdd + nomad_hdd + rrm_hdd + onlineeditor_hdd + monitoring_hdd + logstash_hdd + ario_hdd + dtes_hdd
        +  elasticsearch_hdd + lk_hdd
        )
    #Разделы индексов полнотекстового поиска
    def calculate_serachindex_size(elasticsearch, redundancy, database, main_storage_doc, main_storage_db):
        if elasticsearch.lower() != "false":
            if redundancy.lower() == "true" and database.lower() == "postgres":
                value = main_storage_doc * 0.05 + (main_storage_db * 0.05) / 2
            else:
                value = main_storage_doc * 0.05 + main_storage_db * 0.05
            # Округление вверх до ближайшего целого числа
            result = math.ceil(value)
        else:
            result = 0
        return result
    elasticsearch_serachindex_size = calculate_serachindex_size(elasticsearch, redundancy, database, main_storage_doc, main_storage_db)
    #Разделы индексов системы мониторинга
    if monitoring.lower == "true":
        monitoring_index_size = math.ceil(concurrent_users/100*30)
    else:
        monitoring_index_size = 0
    #Разделы средненагруженных данных (ФХ тел документов) = main storage doc
    #Разделы сервисных баз данных СУБД
    service_db_size = math.ceil(concurrent_users/500*2)
    #Разделы низконагруженных данных (резервное хранение/копирование)
    lowspeed_storage = (reserve_storage_db + reserve_storage_doc)


    #Функция для замены текста в шаблоне
    def replace_placeholder(doc, placeholder, value):
        # Обработка параграфов
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                # Объединение всех runs в одном тексте
                inline = paragraph.runs
                full_text = ''.join([run.text for run in inline])
                if placeholder in full_text:
                    new_text = full_text.replace(placeholder, value)
                    # Очистка существующих runs
                    for run in inline:
                        run.text = ''
                    # Добавление нового текста в первый run
                    inline[0].text = new_text

        # Обработка таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_placeholder(cell, placeholder, value)

    # Заменяем необходимые поля
    replacements = {
        # Блок с общей информацией 
        "CompanyName": str(organization),
        "CurrentDate": str(current_date),
        "UsersPeak": str(concurrent_users),
        "TotalUsers": str(registeredUsers),
        "ImportPeriod": str(dcsdochours),
        "ExtIntegration": str(integrationsystems),
        # Условия фнукционирования
        "OSTypeSQL": str(operationsystem),
        "DBTypeSQL": str(database),
        #Активность пользователей
        "UserCount": str(registeredUsers),
        "UsersForecast": str(concurrent_users),
        "K8SCOUNT": str(k8s_count),
        "K8SCPU": str(k8s_cpu),
        "K8SRAM": str(k8s_ram),
        "K8SHDD": str(k8s_hdd),
        #Блок Веб-сервер
        "WEBCOUNT": str(webserver_count),
        "WEBCPU": str(webserver_cpu),
        "WEBRAM": str(webserver_ram),
        "WEBHDD": str(webserver_hdd),
        #Блок микросервисов
        "MSCOUNT": str(ms_count),
        "MSCPU": str(ms_cpu), 
        "MSRAM": str(ms_ram),
        "MSHDD": str(ms_hdd),
        #Nomad
        "NOMADCOUNT": str(nomad_count),
        "NOMADCPU": str(nomad_cpu),
        "NOMADRAM": str(nomad_ram),
        "NOMADHDD": str(nomad_hdd),
        #ReversePorxy
        "RPCOUNT": str(reverseproxy_count),
        "RPCPU": str(reverseproxy_cpu),
        "RPRAM": str(reverseproxy_ram),
        "RPHDD": str(reverseproxy_hdd),
        #СУБД
        "SQLCOUNT": str(sql_count),
        "SQLCPU": str(sql_cpu),
        "SQLRAM": str(sql_ram),
        "SQLHDD": str("100"),
        #СУБД
        "DCSCOUNT": str(dcs_count),
        "DCSCPU": str(dcs_cpu),
        "DCSRAM": str(dcs_ram),
        "DCSHDD": str(dcs_hdd),
        #Полнотекстовый поиск
        "ELASTICCOUNT": str(elasticsearch_count),
        "ELASTICCPU": str(elasticsearch_cpu),
        "ELASTICRAM": str(elasticsearch_ram),
        "ELASTICHDD": str(elasticsearch_hdd),
        #Мониторинг
        "MONITORINGCOUNT": str(monitoring_count),
        "MONITORINGCPU": str(monitoring_cpu),
        "MONITORINGRAM": str(monitoring_ram),
        "MONITORINGHDD": str(monitoring_hdd),
        #Доп узел Logstash
        "LOGSTASHCOUNT": str(logstash_count),
        "LOGSTASHCPU": str(logstash_cpu),
        "LOGSTASHRAM": str(logstash_ram),
        "LOGSTASHHDD": str(logstash_hdd),
        #Интеграция с онлайн редакторами
        "ONLINEEDITORCOUN": str(onlineeditor_count),
        "ONLINEEDITORCPU": str(onlineeditor_cpu),
        "ONLINEEDITORRAM": str(onlineeditor_ram),
        "ONLINEEDITORHDD": str(onlineeditor_hdd),
        #Узлы АРИО
        "ARIOCOUNT": str(ario_count),
        "ARIOCPU": str(ario_cpu),
        "ARIORAM": str(ario_ram),
        "ARIOHDD": str(ario_hdd),
        "DTESCOUNT": str(dtes_count),
        "DTESCPU": str(dtes_cpu),
        "DTESRAM": str(dtes_ram),
        "DTESHDD": str(dtes_hdd),
        #Узлы RabbitMQ, etcd + keepalived + haproxy 
        "RRMCOUNT": str(rrm_count),
        "RRMCPU": str(rrm_cpu),
        "RRMRAM": str(rrm_ram),
        "RRMHDD": str(rrm_hdd),
        #Сумма ресурсов
        "UnitsCPU": str((webserver_count*webserver_cpu)+(ms_count*ms_cpu)+(k8s_count*k8s_cpu)+(nomad_count*nomad_cpu)+(reverseproxy_count*reverseproxy_cpu)+(sql_count*sql_cpu)
            +(dcs_count*dcs_cpu)+(elasticsearch_count*elasticsearch_cpu)+(monitoring_count*monitoring_cpu)+(ario_count*ario_cpu)+(dtes_count*dtes_cpu)
            +(onlineeditor_count*onlineeditor_cpu)+(lk_count*lk_cpu)
            ), 
        "UnitsRAM": str((webserver_count*webserver_ram)+(ms_count*ms_ram)+(k8s_count*k8s_ram)+(nomad_count*nomad_ram)+(reverseproxy_count*reverseproxy_ram)+(sql_count*sql_ram)
            +(dcs_count*dcs_ram)+(elasticsearch_count*elasticsearch_ram)+(monitoring_count*monitoring_ram)+(ario_count*ario_ram)+(dtes_count*dtes_ram)
            +(onlineeditor_count*onlineeditor_ram)+(lk_count*lk_ram)),
        # Прирост и миграция
        "ImportDataSize": str(importhistorydata_size),
        "YearlyDataSize": str(annualdatagrowth_size),
        "SQLStorageSize": str(main_storage_db),
        "SQLResStorageSize": str(reserve_storage_db),
        "FastStorageSize": str(highspeed_storage),
        "SearchIndexSize": str(elasticsearch_serachindex_size),
        "MidStorageSize": str(main_storage_doc),
        "ServiceDBStorageSize": str(service_db_size),
        "SlowStorageSize": str(lowspeed_storage),
        "FStorageSize": str(main_storage_doc),
        "FResStorageSize": str(reserve_storage_doc),
    }

    for placeholder, value in replacements.items():
        replace_placeholder(doc, placeholder, value)

    # Сохраняем готовый отчет
    report_filename = f"report_{uuid.uuid4()}.docx"
    report_path = os.path.join(app.config['REPORT_FOLDER'], report_filename)
    doc.save(report_path)

    # Удаляем загруженный XML файл (опционально)
    os.remove(filepath)

    report_link = url_for('download_report', filename=report_filename)

    return render_template('index.html', report_link=report_link)

@app.route('/reports/<filename>')
def download_report(filename):
    return send_from_directory(app.config['REPORT_FOLDER'], filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
